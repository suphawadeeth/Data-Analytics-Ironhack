import xlwings as xs

# tell python to open excel name "Book"
new_book = xw.Book()

# open new sheet name "inputs" in the excel name "Book"
inputs = new_book.sheets.add("inputs")

# point to the first sheet in the file and assign into worksheet name
worksheet = new_book.sheets[0]

# write text into the worksheet at the cell A1 using .range().value =
worksheet.range("A1").value = "Hello, how are you?"

# format the cell
worksheet.range("A1").color = (0,100,0)
worksheet.range("A1").column_width = 50

# clear the contents in the cell
worksheet.range("A1").clear_contents()

# clear contents and format
worksheet.range("A1").clear()


# to read data in the cell
name = worksheet.range("C3").value
age = worksheet.range("C4").value
hobby = worksheet.range("C5").value
print(name)
print(age)
print(hobby)

### never referencing things explicitely
### because if we move data to other cell
### you won't be able to read data using three syntaxes above
# instead you can call the cell by its label
name = worksheet.range("name").value
age = worksheet.range("age").value
hobby = worksheet.range("hobby").value

### next example is to combine 2 files together
import pandas as pd
orders = pd.read_excel("filename.xlsx", sheet_name = "Orders")
returns = pd.read_excel("filename.xlsx", sheet_name = "Returns")

# merge 2 files together using the "Order ID" as a link
returned_orders = orders.merge(returns, on ="Order ID")

# calculate total sale of each category (in this return customers group)
# call function .sum()
returned_categories = returned_orders.groupby("category")["Sales", "Profit"].sum()

# visualize it to create a report
import matplotlib.pyplot as plt
fig = plt.figure()
plt.bar(x = returned_categories.index, height = returned_categories["Sale"]

# then put this graph into excel file
# call the file
book = xw.Book("filename.xlxs")

book.sheet #call this syntax to see all the worksheet inside the file

# then select the worksheet you want to put a graph on
book.sheet.add("graph")
target_sheet = book.sheets[0]

target_sheet.picture.add(fig, name = "myplot")

# save a plot into a picture into a png file
plt.savefig("analysis.png")


### Then we can generate a Word report
from docx import Document

#create a new Word document object in python
document = Document()

document.add_heading("top secret report, highly classified", level = 0)

document.add_heading("Executive Summary", level = 1)

document.add_paragraph("The figure below shows the values of returned sales by product category")

document.add_picture("analysis.png")
# save into a file name "amazing_report"
document.save("amazing_report.docx")


### then send email

# this special package works only on windows
import win32com.client as win32

outlook = win32.Despatch("outlook.application")

# createing a new email object
email = outlook.CreateItem(0)

email.To = "emailname.xx@ironhack.com"
email.Subject = "This is a serious report"
email.Body = """Dear all,
                Please find attached ...."""


# last step send with attachments
email.Attachments.Add(Source="here put the path of your file")
email.Send()

##############
##############
### for mac that we don't have outlook
### we can send email using gmail for example

### ## first!
# go to the email account which you want to use for sending email
# then go http://myaccount.google.com/lessercureapps
import smtplib
server = smtplib.SMTP("smtp.gmail.com:587")
server.ehlo()
server.starttls()
server.login("peterramironkack@gmail.com", "unicorns123#") ###  login (usernam, passowrd)

#send the mail
msg = "Hello!"

# sendmail(from which, to which, with msg)
server.sendmail("peterramironkack@gmail.com", "peterramironkack@gmail.com", msg)

server.quit()
